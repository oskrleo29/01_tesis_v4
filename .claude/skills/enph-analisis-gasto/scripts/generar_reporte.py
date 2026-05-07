"""
Generador de reportes .docx en el estilo de la tesis
"ANALISIS DEL GASTO DE LOS HOGARES_*.docx".

Uso típico:
    from generar_reporte import (
        nuevo_documento, agregar_seccion_crispdm,
        tabla_regresion, tabla_clusters, insertar_figura,
        respuesta_jurado_doc
    )

Estilos compatibles con python-docx ≥ 0.8.
"""

from __future__ import annotations
from pathlib import Path
import io
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# Documento base
# ============================================================
def nuevo_documento(titulo: str = "Análisis complementario - Tesis",
                    autora: str = "Paola Velandia") -> Document:
    """Crea un Document con tipografía y márgenes coherentes con la tesis."""
    doc = Document()
    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Tipografía base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Título
    h = doc.add_heading(titulo, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(autora).italic = True
    return doc


def agregar_seccion(doc: Document, titulo: str, nivel: int = 1,
                    intro: str | None = None) -> None:
    """Agrega un encabezado y opcionalmente un párrafo introductorio."""
    doc.add_heading(titulo, level=nivel)
    if intro:
        doc.add_paragraph(intro)


def agregar_seccion_crispdm(doc: Document, fase: str, titulo: str,
                             contenido: str | None = None) -> None:
    """
    Agrega un capítulo correspondiente a una de las seis fases de CRISP-DM.
    `fase` es uno de: 'negocio', 'datos', 'preparacion', 'modelado',
    'evaluacion', 'implementacion'.
    """
    mapa = {
        'negocio': 'Comprensión del negocio',
        'datos': 'Comprensión de los datos',
        'preparacion': 'Preparación de los datos',
        'modelado': 'Fase de modelado',
        'evaluacion': 'Evaluación del modelado',
        'implementacion': 'Implementación y uso operativo',
    }
    encabezado = f"{mapa.get(fase, fase)}: {titulo}"
    doc.add_heading(encabezado, level=1)
    if contenido:
        doc.add_paragraph(contenido)


# ============================================================
# Tablas académicas
# ============================================================
def tabla_regresion(doc: Document,
                    coeficientes: pd.DataFrame,
                    titulo: str = "Tabla X: Coeficientes del modelo WLS ponderado",
                    n_muestra: int | None = None,
                    n_expandido: float | None = None,
                    r2: float | None = None,
                    rmse: float | None = None,
                    cov_type: str = 'HC3') -> None:
    """
    Inserta una tabla de coeficientes con SE entre paréntesis y asteriscos
    de significancia, en formato académico.

    `coeficientes` debe ser un DataFrame con columnas:
        - 'variable'
        - 'coef'
        - 'se'
        - 'p_valor'
    """
    p = doc.add_paragraph()
    p.add_run(titulo).bold = True

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Variable'
    hdr[1].text = 'Coeficiente (EE)'

    for _, fila in coeficientes.iterrows():
        sig = ''
        p_val = fila['p_valor']
        if pd.notna(p_val):
            if p_val < 0.01:
                sig = '***'
            elif p_val < 0.05:
                sig = '**'
            elif p_val < 0.10:
                sig = '*'
        celda = f"{fila['coef']:.4f}{sig}\n({fila['se']:.4f})"
        row = table.add_row().cells
        row[0].text = str(fila['variable'])
        row[1].text = celda

    # Pie de tabla
    nota_partes = []
    if n_muestra is not None:
        nota_partes.append(f"N muestral = {n_muestra:,}")
    if n_expandido is not None:
        nota_partes.append(f"N expandido = {n_expandido:,.0f}")
    if r2 is not None:
        nota_partes.append(f"R² = {r2:.4f}")
    if rmse is not None:
        nota_partes.append(f"RMSE = {rmse:.4f}")
    nota_base = (f"Nota: errores estándar entre paréntesis, robustos a heterocedasticidad "
                 f"({cov_type}). *** p<0.01, ** p<0.05, * p<0.10. "
                 f"Variable dependiente: log_gastos_2025. Pesos: FEX_C normalizado.")
    if nota_partes:
        nota_base += " " + ". ".join(nota_partes) + "."
    nota = doc.add_paragraph()
    r = nota.add_run(nota_base)
    r.italic = True
    r.font.size = Pt(9)


def tabla_clusters(doc: Document,
                   tabla: pd.DataFrame,
                   titulo: str = "Tabla X: Caracterización de clusters K-Means ponderado por FEX_C") -> None:
    """
    Inserta una tabla de clusters en el formato típico de la tesis.

    `tabla` debe ser un DataFrame de `caracterizar_clusters_ponderado()` con:
    cluster_ord_w, n_muestra, n_expandido, gasto_log_*, ingreso_log_*, part_pct_expandido.
    """
    p = doc.add_paragraph()
    p.add_run(titulo).bold = True

    cols = ['cluster_ord_w', 'n_muestra', 'n_expandido', 'part_pct_expandido',
            'gasto_log_mediana', 'ingreso_log_mediana']
    cols = [c for c in cols if c in tabla.columns]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Light Grid Accent 1'
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = {
            'cluster_ord_w': 'Cluster',
            'n_muestra': 'N muestral',
            'n_expandido': 'N expandido',
            'part_pct_expandido': '% expandido',
            'gasto_log_mediana': 'log gasto mediana',
            'ingreso_log_mediana': 'log ingreso mediana',
        }.get(c, c)

    for _, fila in tabla.iterrows():
        row = table.add_row().cells
        for i, c in enumerate(cols):
            val = fila[c]
            if isinstance(val, (int, np.integer)):
                row[i].text = f"{val:,}"
            elif c == 'n_expandido':
                row[i].text = f"{val:,.0f}"
            elif c == 'part_pct_expandido':
                row[i].text = f"{val:.1f}%"
            elif isinstance(val, (float, np.floating)):
                row[i].text = f"{val:.4f}"
            else:
                row[i].text = str(val)

    nota = doc.add_paragraph()
    r = nota.add_run(
        "Nota: clusters ordenados por ingreso mediano (0 = menor, 3 = mayor). "
        "Pesos de expansión normalizados por la media (FEX_C). "
        "Fuente: cálculos propios con base en ENPH 2016-2017, DANE."
    )
    r.italic = True
    r.font.size = Pt(9)


def tabla_ratio_percentil(doc: Document,
                           ratio_table: dict,
                           titulo: str = "Tabla X: Percentiles del ratio gasto/ingreso por cluster") -> None:
    """
    Inserta la RATIO_TABLE_W como tabla en el .docx.
    `ratio_table` es el dict producido en celda 144 del notebook.
    """
    p = doc.add_paragraph()
    p.add_run(titulo).bold = True

    cols = ['cluster', 'p25', 'p50', 'p75', 'p90']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Light Grid Accent 1'
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = c.upper()

    for cluster_id, percs in sorted(ratio_table.items()):
        row = table.add_row().cells
        row[0].text = str(cluster_id)
        for i, q in enumerate(['p25', 'p50', 'p75', 'p90'], start=1):
            row[i].text = f"{percs.get(q, np.nan):.4f}"

    nota = doc.add_paragraph()
    r = nota.add_run(
        "Nota: ratio = exp(log_gastos_2025) / exp(log_ingresos_2025). "
        "Valores ponderados por FEX_C."
    )
    r.italic = True
    r.font.size = Pt(9)


# ============================================================
# Figuras
# ============================================================
def insertar_figura(doc: Document,
                    fig,
                    titulo: str,
                    nota: str | None = None,
                    width_inches: float = 6.0,
                    dpi: int = 300) -> None:
    """
    Inserta una figura matplotlib (`fig`) en el documento con título y nota.
    """
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
    buf.seek(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(titulo)
    r.bold = True
    r.font.size = Pt(10)

    if nota:
        n = doc.add_paragraph()
        n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = n.add_run(nota)
        rn.italic = True
        rn.font.size = Pt(9)


# ============================================================
# Documento de respuesta al jurado
# ============================================================
def respuesta_jurado_doc(observaciones: list[dict],
                          titulo_tesis: str = "Estimación del gasto personal "
                                              "como variable principal para "
                                              "la evaluación de capacidad de pago",
                          autora: str = "Paola Velandia",
                          jurado: str = "") -> Document:
    """
    Crea un .docx en el formato de carta de respuesta al jurado.

    `observaciones` es una lista de dicts, cada uno con:
        - 'observacion': texto literal de la observación.
        - 'respuesta': texto de la respuesta.
        - 'ubicacion' (opcional): página/sección donde se ve el cambio.
        - 'tabla' (opcional): DataFrame con coef si se quiere insertar.
    """
    doc = Document()

    h = doc.add_heading('Respuesta a observaciones del jurado evaluador', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(f'Tesis: {titulo_tesis}\n').italic = True
    p.add_run(f'Autora: {autora}\n').italic = True
    if jurado:
        p.add_run(f'Jurado: {jurado}\n').italic = True

    intro = doc.add_paragraph(
        "Se agradecen las observaciones detalladas a la versión preliminar "
        "de la tesis. A continuación se responde a cada punto. Las "
        "modificaciones realizadas en el manuscrito se identifican con "
        "marcas en la versión revisada que se anexa."
    )

    for i, obs in enumerate(observaciones, start=1):
        doc.add_heading(f'Observación {i}', level=2)

        p = doc.add_paragraph()
        r = p.add_run(obs.get('observacion', ''))
        r.italic = True

        p = doc.add_paragraph()
        p.add_run('Respuesta: ').bold = True
        p.add_run(obs.get('respuesta', ''))

        if 'ubicacion' in obs and obs['ubicacion']:
            p = doc.add_paragraph()
            r = p.add_run(f'Ubicación del cambio: {obs["ubicacion"]}')
            r.italic = True
            r.font.size = Pt(9)

        if 'tabla' in obs and obs['tabla'] is not None:
            df = obs['tabla']
            t = doc.add_table(rows=1, cols=len(df.columns))
            t.style = 'Light Grid Accent 1'
            for j, c in enumerate(df.columns):
                t.rows[0].cells[j].text = str(c)
            for _, fila in df.iterrows():
                row = t.add_row().cells
                for j, c in enumerate(df.columns):
                    val = fila[c]
                    if isinstance(val, (float, np.floating)):
                        row[j].text = f"{val:.4f}"
                    else:
                        row[j].text = str(val)

    return doc


# ============================================================
# Helpers de exportación
# ============================================================
def guardar(doc: Document, ruta: str | Path) -> Path:
    p = Path(ruta)
    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return p


if __name__ == '__main__':
    print("Módulo generar_reporte.py")
    print("Funciones disponibles:")
    for f in [nuevo_documento, agregar_seccion, agregar_seccion_crispdm,
              tabla_regresion, tabla_clusters, tabla_ratio_percentil,
              insertar_figura, respuesta_jurado_doc, guardar]:
        print(f"  - {f.__name__}")
